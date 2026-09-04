"""Generate client-facing Word document for store account setup."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT = "UGolf-Store-Account-Setup-Guide.docx"


def set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(text)


def main():
    doc = Document()
    set_normal_style(doc)

    # Cover
    doc.add_paragraph()
    doc.add_paragraph()
    add_title(doc, "UGolf Mobile App")
    add_subtitle(doc, "Google Play & Apple App Store\nDeveloper Account Setup Guide")
    doc.add_paragraph()
    add_subtitle(doc, "Prepared for: Client / App Owner")
    add_subtitle(doc, "Document version: 1.0  |  Date: August 2026")
    doc.add_page_break()

    # Section 1
    add_heading(doc, "1. Purpose of This Document")
    add_para(
        doc,
        "This guide explains how to create and configure the developer accounts required "
        "to publish the UGolf mobile app on the Google Play Store (Android) and the "
        "Apple App Store (iOS). Please follow the steps in this document and share the "
        "requested access details with your development team once setup is complete.",
    )

    # Section 2
    add_heading(doc, "2. Overview")
    add_para(doc, "To publish UGolf, you need two separate developer accounts:")
    add_table(
        doc,
        ["Platform", "Account Name", "Cost", "Renewal"],
        [
            ["Android", "Google Play Developer Account", "USD 25 (one-time)", "No renewal fee"],
            ["iOS", "Apple Developer Program", "USD 99 per year", "Annual renewal required"],
        ],
    )
    add_para(
        doc,
        "These accounts must be owned by the party that will legally publish and "
        "operate the app (you or your company). The development team will use these "
        "accounts to upload builds, but the account owner retains full control.",
    )

    # Section 3
    add_heading(doc, "3. Individual vs Organization Account")
    add_para(
        doc,
        "Both Google and Apple allow you to register as an Individual or an Organization. "
        "Choose the option that matches how you want the app to appear on the stores.",
    )
    add_table(
        doc,
        ["Account Type", "Store Listing Shows", "Best For"],
        [
            [
                "Individual",
                "Your personal name",
                "Solo founders, personal projects, early-stage launches",
            ],
            [
                "Organization",
                "Your company name (e.g. UGolf Pvt Ltd)",
                "Registered businesses, teams, client-facing brands",
            ],
        ],
    )
    add_note(
        doc,
        "If UGolf is a company product, an Organization account is recommended so the "
        "app is published under your business name. Organization accounts require "
        "additional business verification (including a D-U-N-S number for Apple).",
    )

    # Section 4
    add_heading(doc, "4. Information to Prepare Before You Start")
    add_para(doc, "Gather the following before beginning registration:")
    add_bullets(
        doc,
        [
            "Valid government-issued photo ID (passport or national ID)",
            "Business email address (recommended: dev@yourcompany.com or similar)",
            "Company legal name and registered address (for Organization accounts)",
            "D-U-N-S Number (required for Apple Organization accounts — free from Dun & Bradstreet)",
            "Payment card for registration fees (Visa, MasterCard, or equivalent)",
            "Privacy Policy URL (a public webpage describing how user data is collected and used)",
            "Support email and website URL for the app",
            "App icon (1024 x 1024 px) and store screenshots (development team can assist)",
        ],
    )

    # Section 5 - Google Play
    add_heading(doc, "5. Google Play Developer Account Setup (Android)")
    add_heading(doc, "Step 1: Create or Use a Google Account", level=2)
    add_para(
        doc,
        "Use a dedicated Google account for app publishing. Avoid using a personal account "
        "that may be deactivated or lost. A business email linked to Google Workspace is ideal.",
    )

    add_heading(doc, "Step 2: Register as a Google Play Developer", level=2)
    add_numbered(
        doc,
        [
            "Open https://play.google.com/console/signup in your browser.",
            "Sign in with your Google account.",
            "Read and accept the Google Play Developer Distribution Agreement.",
            "Pay the one-time registration fee of USD 25.",
            "Complete identity verification when prompted (upload ID; approval may take 1–3 business days).",
        ],
    )

    add_heading(doc, "Step 3: Complete Account Profile", level=2)
    add_numbered(
        doc,
        [
            "Go to Play Console → Settings → Developer account.",
            "Enter your developer name (this appears on the store listing).",
            "Add contact email, phone number, and physical address.",
            "Select account type: Individual or Organization.",
            "If Organization: provide business name, registration details, and complete verification.",
        ],
    )

    add_heading(doc, "Step 4: Set Up Payments Profile (If Applicable)", level=2)
    add_para(
        doc,
        "Required only if the app will be paid or include in-app purchases. "
        "Go to Play Console → Monetize → Monetization setup and link a Google Payments merchant profile.",
    )

    add_heading(doc, "Step 5: Invite Your Development Team", level=2)
    add_numbered(
        doc,
        [
            "Go to Play Console → Users and permissions.",
            "Click Invite new users.",
            "Enter the developer team email addresses provided by your agency.",
            "Assign appropriate roles (typically Admin or Release manager).",
            "Click Send invitation.",
        ],
    )

    add_heading(doc, "Step 6: Create the UGolf App Entry", level=2)
    add_numbered(
        doc,
        [
            "In Play Console, click All apps → Create app.",
            "App name: UGolf",
            "Default language: English (or your preferred language)",
            "App or game: App",
            "Free or paid: Free (unless you plan to charge upfront)",
            "Confirm declarations and create the app.",
        ],
    )
    add_note(
        doc,
        "Your development team will handle uploading the Android build (.aab file), "
        "completing store listing content, and submitting for review once they have access.",
    )

    doc.add_page_break()

    # Section 6 - Apple
    add_heading(doc, "6. Apple Developer Program Setup (iOS)")
    add_heading(doc, "Prerequisites", level=2)
    add_bullets(
        doc,
        [
            "A Mac computer is required for the development team to build and upload iOS releases.",
            "An Apple ID (use a business email, not a shared personal account).",
            "For Organization accounts: D-U-N-S Number and legal entity documents.",
        ],
    )

    add_heading(doc, "Step 1: Create an Apple ID (If Needed)", level=2)
    add_numbered(
        doc,
        [
            "Go to https://appleid.apple.com and create an Apple ID using your business email.",
            "Enable two-factor authentication (required for developer enrollment).",
        ],
    )

    add_heading(doc, "Step 2: Enroll in the Apple Developer Program", level=2)
    add_numbered(
        doc,
        [
            "Go to https://developer.apple.com/programs/enroll/",
            "Sign in with your Apple ID.",
            "Choose enrollment type: Individual or Organization.",
            "Complete the enrollment form with legal name, address, and contact details.",
            "Pay the annual fee of USD 99.",
            "Wait for Apple to verify your identity or organization (Individual: 24–48 hours; Organization: up to several weeks).",
        ],
    )

    add_heading(doc, "Step 3: Access App Store Connect", level=2)
    add_numbered(
        doc,
        [
            "Once enrolled, go to https://appstoreconnect.apple.com",
            "Sign in with the same Apple ID used for enrollment.",
            "Accept any terms and conditions if prompted.",
        ],
    )

    add_heading(doc, "Step 4: Invite Your Development Team", level=2)
    add_numbered(
        doc,
        [
            "In App Store Connect, go to Users and Access.",
            "Click the + button to add a new user.",
            "Enter the developer team email addresses provided by your agency.",
            "Assign roles: Developer (for uploading builds) and/or App Manager (for store listing).",
            "Send the invitation.",
        ],
    )

    add_heading(doc, "Step 5: Create the UGolf App Record", level=2)
    add_numbered(
        doc,
        [
            "In App Store Connect, go to My Apps → + → New App.",
            "Platforms: iOS",
            "Name: UGolf",
            "Primary language: English",
            "Bundle ID: Select or register com.ugolf (your development team will confirm the exact ID).",
            "SKU: ugolf-ios-001 (any unique internal reference).",
            "User Access: Full Access.",
            "Click Create.",
        ],
    )
    add_note(
        doc,
        "Your development team will register the Bundle ID in the Apple Developer portal, "
        "configure signing certificates, upload builds, and complete the App Store listing.",
    )

    doc.add_page_break()

    # Section 7
    add_heading(doc, "7. App Store Requirements Checklist")
    add_para(doc, "The following items must be ready before the app can be submitted for review:")
    add_table(
        doc,
        ["Item", "Description", "Who Provides"],
        [
            ["Privacy Policy URL", "Public webpage explaining data collection (location, account info, etc.)", "Client / Legal"],
            ["Support URL or email", "Contact method for app users", "Client"],
            ["App description", "Short and full description for store listing", "Client + Dev team"],
            ["App icon", "1024 x 1024 px, no transparency (iOS)", "Client / Design team"],
            ["Screenshots", "Phone screenshots for each required device size", "Dev team"],
            ["Demo login (if app requires sign-in)", "Test username and password for store reviewers", "Client"],
            ["Age rating questionnaire", "Completed in both store consoles", "Client + Dev team"],
            ["Data privacy declarations", "Location data usage must be declared (UGolf uses GPS for course maps)", "Client + Dev team"],
        ],
    )

    # Section 8
    add_heading(doc, "8. What to Send Back to Your Development Team")
    add_para(doc, "After completing account setup, please confirm the following with your development team:")
    add_numbered(
        doc,
        [
            "Google Play Console: Account is active and team invitations have been accepted.",
            "Apple Developer Program: Enrollment is approved and team invitations have been accepted.",
            "App Store Connect: UGolf app record has been created.",
            "Privacy Policy URL is live and accessible.",
            "Support email address for the app.",
            "Preferred developer/store listing name (Individual name or Company name).",
            "Demo account credentials for app review (if login is required).",
            "Confirmation of whether the app will be Free, Paid, or include In-App Purchases.",
        ],
    )

    # Section 9
    add_heading(doc, "9. Estimated Timelines")
    add_table(
        doc,
        ["Step", "Estimated Time"],
        [
            ["Google Play registration + ID verification", "1–3 business days"],
            ["Apple Individual enrollment", "24–48 hours"],
            ["Apple Organization enrollment", "1–4 weeks"],
            ["First app review (Google Play)", "3–7 days"],
            ["First app review (Apple App Store)", "1–3 days (can vary)"],
        ],
    )

    # Section 10
    add_heading(doc, "10. Important Notes for UGolf")
    add_bullets(
        doc,
        [
            "UGolf uses location services (GPS) to show the player's position on the course map. This must be declared in both store privacy forms.",
            "A Privacy Policy is mandatory on both platforms before the app can be published.",
            "The Android app package name is: com.ugolf",
            "The iOS bundle identifier will be confirmed and aligned by the development team before submission.",
            "Keep your developer account credentials secure. Do not share passwords — use the built-in invitation/team access features instead.",
            "Apple Developer membership must be renewed annually (USD 99) or your app will be removed from the App Store.",
        ],
    )

    # Section 11
    add_heading(doc, "11. Useful Links")
    add_bullets(
        doc,
        [
            "Google Play Console: https://play.google.com/console",
            "Apple Developer Program: https://developer.apple.com/programs/",
            "App Store Connect: https://appstoreconnect.apple.com",
            "D-U-N-S Number (free): https://www.dnb.com/duns-number/get-a-duns.html",
            "Apple ID: https://appleid.apple.com",
        ],
    )

    # Footer
    doc.add_paragraph()
    add_para(
        doc,
        "If you have questions during setup, contact your development team before completing "
        "payment or creating duplicate app entries.",
        bold=False,
    )

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
